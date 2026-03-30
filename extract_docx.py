import os
from docx import Document
import json
import uuid
from datetime import datetime

target_files = [
    "/home/user/uploaded_files/実習報告書の書式(初等教育実践研究Ａ).第１０回　１１月３０日　後期.docx",
    "/home/user/uploaded_files/実習報告書の書式(初等教育実践研究Ａ).2017　３年　後期　９月２２日　第１回　報告書.docx",
    "/home/user/uploaded_files/実習報告書の書式(初等教育実践研究Ａ).2017 ３年　後期　９月２９日　第２回　報告書.docx",
    "/home/user/uploaded_files/実習報告書の書式(初等教育実践研究Ａ).第１１回　１２月７日　後期　報告書.docx",
    "/home/user/uploaded_files/実習報告書の書式(初等教育実践研究Ａ).第１２回　１２月１４日　後期　報告書.docx"
]

data = []

def extract_text(doc):
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() and cell.text.strip() not in full_text:
                    full_text.append(cell.text.strip())
    return "\n".join(full_text)

for file in target_files:
    if not os.path.exists(file):
        print(f"Not found: {file}")
        continue
    doc = Document(file)
    text = extract_text(doc)
    
    reflection_text = text
    parts = text.split("学ばせていただいた内容")
    if len(parts) > 1:
        reflection_text = "学ばせていただいた内容" + parts[1]
    
    fname = os.path.basename(file)
    week = 1
    if "第１回" in fname: week = 1
    elif "第２回" in fname: week = 2
    elif "第１０回" in fname: week = 10
    elif "第１１回" in fname: week = 11
    elif "第１２回" in fname: week = 12
    
    entry_date = "2023-05-10"
    if "１１月３０日" in fname: entry_date = "2023-11-30"
    elif "９月２２日" in fname: entry_date = "2023-09-22"
    elif "９月２９日" in fname: entry_date = "2023-09-29"
    elif "１２月７日" in fname: entry_date = "2023-12-07"
    elif "１２月１４日" in fname: entry_date = "2023-12-14"

    entry = {
        "id": "journal-" + str(uuid.uuid4())[:8],
        "student_id": "user-001", # テスト用学生ユーザー
        "title": f"第{week}回 実習日誌",
        "entry_date": entry_date,
        "week_number": week,
        "content": text,
        "reflection_text": reflection_text,
        "status": "submitted",
        "created_at": datetime.now().isoformat()
    }
    data.append(entry)

with open("/home/user/webapp/import_data.json", "w", encoding="utf-8") as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print(f"Extracted {len(data)} documents.")
