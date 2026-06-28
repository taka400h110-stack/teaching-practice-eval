#!/usr/bin/env python3
"""Convert the 6-factor rubric Markdown to a formatted .docx (python-docx).

Handles: # / ## / ### headings, | tables |, > blockquotes, ``` code blocks,
horizontal rules (---), and inline **bold**. Tailored to the rubric document
structure (not a general-purpose MD converter).
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_inline(paragraph, text):
    """Add text with **bold** segments."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def parse_table(lines, start):
    """Parse a markdown table starting at index `start`. Return (rows, next_idx)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    # rows[1] is the separator (---). Drop it.
    parsed = []
    for r_i, raw in enumerate(rows):
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        parsed.append(cells)
    if len(parsed) >= 2:
        # remove separator row (index 1)
        body = [parsed[0]] + parsed[2:]
    else:
        body = parsed
    return body, i


def main(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            i += 1
            continue

        # code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            para = doc.add_paragraph()
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            continue

        # headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            add_inline(para, text)
            for run in para.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # table
        if stripped.startswith("|"):
            body, next_i = parse_table(lines, i)
            if body:
                cols = max(len(r) for r in body)
                table = doc.add_table(rows=0, cols=cols)
                table.style = "Light Grid Accent 1"
                for r_idx, row_cells in enumerate(body):
                    cells = row_cells + [""] * (cols - len(row_cells))
                    row = table.add_row()
                    for c_idx, cell_text in enumerate(cells):
                        cell = row.cells[c_idx]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_inline(para, cell_text)
                        if r_idx == 0:
                            for run in para.runs:
                                run.bold = True
            i = next_i
            continue

        # normal paragraph (strip trailing markdown hard-break spaces)
        para = doc.add_paragraph()
        add_inline(para, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Saved {docx_path}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
